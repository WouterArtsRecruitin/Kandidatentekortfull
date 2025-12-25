"""
File text extraction for PDF and DOCX files.
"""

import io
import requests
from typing import Optional
from .logging_config import get_logger
from .retry import retry_with_backoff
from ..config import TYPEFORM_API_TOKEN

logger = get_logger("file_extractor")

# Check for optional dependencies
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed - PDF extraction disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed - DOCX extraction disabled")


@retry_with_backoff(max_attempts=2, initial_delay=1.0, exceptions=(requests.RequestException,))
def download_file(file_url: str) -> bytes:
    """Download file with authentication if needed."""
    headers = {}

    # Typeform files require Bearer token
    if TYPEFORM_API_TOKEN and 'typeform.com' in file_url:
        headers['Authorization'] = f'Bearer {TYPEFORM_API_TOKEN}'
        logger.info("Using Typeform API authentication")

    response = requests.get(file_url, headers=headers, timeout=30)
    response.raise_for_status()

    logger.info(f"Downloaded file: {len(response.content)} bytes, type={response.headers.get('content-type', 'unknown')}")
    return response.content


def detect_file_type(content: bytes, file_url: str = "") -> str:
    """Detect file type from magic bytes or URL."""
    # Check magic bytes first (most reliable)
    if content[:4] == b'%PDF':
        return 'pdf'
    elif content[:2] == b'PK':  # ZIP-based formats (DOCX, XLSX)
        return 'docx'
    elif content[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':  # OLE format (old .doc)
        return 'doc_legacy'

    # Fallback to URL extension
    file_url_lower = file_url.lower()
    if '.pdf' in file_url_lower:
        return 'pdf'
    elif '.docx' in file_url_lower:
        return 'docx'
    elif '.doc' in file_url_lower:
        return 'doc_legacy'

    return 'unknown'


def extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF content."""
    if not PDF_AVAILABLE:
        logger.error("PyPDF2 not available")
        return ""

    try:
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts)
        logger.info(f"PDF extracted: {len(full_text)} chars from {len(reader.pages)} pages")
        return full_text.strip()

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX content."""
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available")
        return ""

    try:
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        logger.info(f"DOCX extracted: {len(full_text)} chars")
        return full_text.strip()

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text_from_file(file_url: str) -> Optional[str]:
    """
    Download and extract text from a file URL.
    Supports PDF and DOCX formats.

    Returns:
        Extracted text or empty string on failure
    """
    if not file_url:
        return ""

    try:
        logger.info(f"Extracting text from: {file_url[:80]}...")

        # Download file
        content = download_file(file_url)

        # Check for error responses
        if len(content) < 100 and b'error' in content.lower():
            logger.error(f"Got error response: {content[:200]}")
            return ""

        # Detect and extract
        file_type = detect_file_type(content, file_url)

        if file_type == 'pdf':
            return extract_pdf_text(content)
        elif file_type == 'docx':
            return extract_docx_text(content)
        elif file_type == 'doc_legacy':
            logger.warning("Legacy .doc format not supported - please use .docx")
            return ""
        else:
            logger.warning(f"Unknown file type. First bytes: {content[:20]}")
            return ""

    except Exception as e:
        logger.error(f"File extraction error: {e}")
        return ""
